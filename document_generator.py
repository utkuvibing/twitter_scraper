"""Export archived X posts to supported document formats."""

import csv
import io
import logging
import os
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from export_schema import (
    EXPORT_SCHEMA_VERSION,
    atomic_save_docx,
    atomic_write_json,
    atomic_write_text,
    build_export_payload,
    csv_safe,
    normalize_tweet,
    resolve_output_path,
)
from time_utils import utc_now

logger = logging.getLogger(__name__)


def default_output_dir() -> str:
    """Return the predictable writable output root for the current invocation."""
    return str((Path.cwd() / "output").resolve())


BASE_OUTPUT_DIR = default_output_dir()


def _markdown_quote(value: str) -> str:
    return "\n".join(f"> {line}" for line in value.splitlines())


CSV_COLUMNS = (
    "id",
    "date",
    "date_str",
    "text",
    "tweet_url",
    "has_media",
    "media_urls",
    "has_article",
    "needs_full_text",
)


def get_output_path(target_username: str, filename: str, output_dir: Optional[str] = None) -> str:
    """
    Kullanıcıya özel output klasörü oluştur ve tam yol döndür

    Args:
        target_username: Hedef kullanıcı adı
        filename: Dosya adı
        output_dir: Özel output dizini (None ise default output kullanılır)

    Returns:
        Tam dosya yolu
    """
    _, ext = os.path.splitext(filename)
    return resolve_output_path(
        target_username=target_username,
        filename=filename,
        extension=ext or ".json",
        base_output_dir=default_output_dir(),
        output_dir=output_dir,
    )


def create_word_document(
    tweets: List,
    output_path: str,
    target_username: str,
    output_dir: Optional[str] = None,
) -> str:
    """
    Tweetleri Word document olarak kaydet

    Args:
        tweets: Tweet listesi (scraper.Tweet objeleri)
        output_path: Çıktı dosya yolu (.docx)
        target_username: Scrape edilen hesabın kullanıcı adı
        output_dir: Özel output dizini (None ise default output kullanılır)

    Returns:
        Kaydedilen dosya yolu
    """
    doc = Document()

    # Başlık stili
    title = doc.add_heading(f"@{target_username} - Post Archive", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta bilgi
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Total: {len(tweets)} posts | Exported: {utc_now().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Boşluk

    # Tweetleri ekle (güncel'den eskiye - zaten bu sırada)
    for i, tweet in enumerate(tweets, 1):
        # Tweet numarası ve tarih
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(f"Post #{i}")
        header_run.bold = True
        header_run.font.size = Pt(12)
        header_run.font.color.rgb = RGBColor(29, 155, 240)  # X mavi rengi

        date_run = header_para.add_run(f"  |  {tweet.date_str}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(100, 100, 100)

        # Tweet metni
        if tweet.text:
            text_para = doc.add_paragraph()
            text_para.add_run(tweet.text)

        # Medya linkleri
        if tweet.media_urls:
            media_para = doc.add_paragraph()
            media_label = media_para.add_run("Media: ")
            media_label.bold = True
            media_label.font.size = Pt(9)
            media_label.font.color.rgb = RGBColor(80, 80, 80)

            for j, url in enumerate(tweet.media_urls):
                if j > 0:
                    media_para.add_run(" | ")
                media_link = media_para.add_run(url if len(url) < 80 else url[:77] + "...")
                media_link.font.size = Pt(8)
                media_link.font.color.rgb = RGBColor(100, 100, 100)

        # Tweet URL
        url_para = doc.add_paragraph()
        url_label = url_para.add_run("Link: ")
        url_label.font.size = Pt(9)
        url_label.font.color.rgb = RGBColor(80, 80, 80)
        url_link = url_para.add_run(tweet.tweet_url)
        url_link.font.size = Pt(9)
        url_link.font.color.rgb = RGBColor(29, 155, 240)

        # Ayırıcı çizgi
        separator = doc.add_paragraph()
        sep_run = separator.add_run("─" * 60)
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)

    # Dosyayı kaydet
    if not output_path.endswith(".docx"):
        output_path += ".docx"

    # Kullanıcıya özel output klasörüne kaydet (seçilen dizini kullan)
    full_path = resolve_output_path(
        target_username, output_path, ".docx", default_output_dir(), output_dir
    )
    atomic_save_docx(doc, full_path)
    logger.info("Word document saved: %s", full_path)
    print(f"Word document saved: {full_path}")

    return full_path


def create_simple_document(
    tweets: List,
    output_path: str,
    target_username: str,
    output_dir: Optional[str] = None,
) -> str:
    """
    Basit formatta Word document oluştur (sadece metin ve tarih)

    Args:
        tweets: Tweet listesi
        output_path: Çıktı dosya yolu
        target_username: Kullanıcı adı
        output_dir: Özel output dizini (None ise default output kullanılır)

    Returns:
        Dosya yolu
    """
    doc = Document()

    doc.add_heading(f"@{target_username} Posts", 0)
    doc.add_paragraph(f"Total: {len(tweets)} posts")
    doc.add_paragraph()

    for i, tweet in enumerate(tweets, 1):
        para = doc.add_paragraph()
        para.add_run(f"[{i}] {tweet.date_str}\n").bold = True
        para.add_run(tweet.text or "[No text - media only]")
        para.add_run("\n")

    if not output_path.endswith(".docx"):
        output_path += ".docx"

    full_path = resolve_output_path(
        target_username, output_path, ".docx", default_output_dir(), output_dir
    )
    atomic_save_docx(doc, full_path)
    return full_path


def create_json_document(
    tweets: List,
    output_path: str,
    target_username: str,
    output_dir: Optional[str] = None,
    scrape_type: str = "profile",
) -> str:
    """
    Tweetleri JSON formatında kaydet (MCP-ready)

    Args:
        tweets: Tweet listesi
        output_path: Çıktı dosya yolu
        target_username: Kullanıcı adı
        output_dir: Özel output dizini (None ise default output kullanılır)
        scrape_type: Export metadata source type ("profile" or "bookmarks")

    Returns:
        Dosya yolu
    """
    data = build_export_payload(tweets, target_username, scrape_type=scrape_type)

    full_path = resolve_output_path(
        target_username, output_path, ".json", default_output_dir(), output_dir
    )
    atomic_write_json(full_path, data)
    logger.info("JSON export saved: %s", full_path)
    print(f"JSON saved: {full_path}")
    return full_path


def create_csv_document(
    tweets: List,
    output_path: str,
    target_username: str,
    output_dir: Optional[str] = None,
) -> str:
    """Write normalized tweets as a spreadsheet-friendly CSV file."""
    stream = io.StringIO(newline="")
    writer = csv.DictWriter(stream, fieldnames=CSV_COLUMNS, lineterminator="\n")
    writer.writeheader()

    for tweet in tweets:
        normalized = normalize_tweet(tweet)
        normalized["media_urls"] = " | ".join(normalized["media_urls"])
        writer.writerow({column: csv_safe(normalized[column]) for column in CSV_COLUMNS})

    full_path = resolve_output_path(
        target_username, output_path, ".csv", default_output_dir(), output_dir
    )
    atomic_write_text(full_path, stream.getvalue(), encoding="utf-8-sig")
    logger.info("CSV saved: %s", full_path)
    print(f"CSV saved: {full_path}")
    return full_path


def create_markdown_document(
    tweets: List,
    output_path: str,
    target_username: str,
    output_dir: Optional[str] = None,
) -> str:
    """
    Tweetleri Markdown formatında kaydet

    Args:
        tweets: Tweet listesi
        output_path: Çıktı dosya yolu
        target_username: Kullanıcı adı
        output_dir: Özel output dizini (None ise default output kullanılır)

    Returns:
        Dosya yolu
    """
    lines = []
    normalized_tweets = [normalize_tweet(tweet) for tweet in tweets]

    lines.append(f"# @{target_username} - Post Archive\n")
    lines.append(f"**Schema:** {EXPORT_SCHEMA_VERSION}\n")
    lines.append(f"**Total:** {len(tweets)} posts\n")
    lines.append(f"**Exported:** {utc_now().strftime('%Y-%m-%d %H:%M UTC')}\n")
    lines.append("---\n")

    for i, tweet in enumerate(normalized_tweets, 1):
        lines.append(f"## Post #{i}\n")
        lines.append(f"**Date:** {tweet['date_str']}\n")
        if tweet["text"]:
            lines.append(f"\n{_markdown_quote(tweet['text'])}\n")
        if tweet["media_urls"]:
            lines.append(f"\n**Media:** {len(tweet['media_urls'])} items\n")
        lines.append(f"\n[Post link]({tweet['tweet_url']})\n")
        lines.append("\n---\n")

    full_path = resolve_output_path(
        target_username, output_path, ".md", default_output_dir(), output_dir
    )
    atomic_write_text(full_path, "\n".join(lines))

    logger.info("Markdown export saved: %s", full_path)
    print(f"Markdown saved: {full_path}")
    return full_path
